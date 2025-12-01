# File: /HttpTriggerGetData/__init__.py
import os
import io
import json
import logging
import requests
import azure.functions as func
from datetime import datetime, timedelta
from docx import Document


def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("Fetching paged data and generating Word document...")

    try:
        # === Environment variables ===
        auth_url = os.environ["API_AUTH_URL"]
        data_url = os.environ["API_DATA_URL"]
        username = os.environ["API_USERNAME"]
        password = os.environ["API_PASSWORD"]

        # === Date range ===
        to_dt = datetime.utcnow()
        from_dt = to_dt - timedelta(days=7)
        from_param = from_dt.strftime("%Y/%m/%d %H:%M:%S.000Z")
        to_param = to_dt.strftime("%Y/%m/%d %H:%M:%S.000Z")

        # === Step 1: Authenticate ===
        auth_headers = {
            "Accept": "application/json",
            "Content-Type": "application/json; charset=utf-8",
        }
        auth_payload = {"Username": username, "Password": password}

        auth_resp = requests.post(auth_url, headers=auth_headers, json=auth_payload, timeout=60)
        auth_resp.raise_for_status()
        auth_data = auth_resp.json()

        token = (
            auth_data.get("Token")
            or auth_data.get("token")
            or auth_data.get("Data", {}).get("Token")
            or auth_data.get("data", {}).get("token")
        )
        if not token:
            raise ValueError(f"No token found in auth response: {auth_data}")

        # === Step 2: Helper to fetch a single page ===
        def fetch_page(page_num: int):
            params = {
                "From": from_param,
                "To": to_param,
                "Page": page_num,
            }
            headers = {"Accept": "application/json", "Authorization": f"Bearer {token}"}

            resp = requests.get(data_url, headers=headers, params=params, timeout=60)
            if resp.status_code == 401:
                headers["Authorization"] = token
                resp = requests.get(data_url, headers=headers, params=params, timeout=60)

            resp.raise_for_status()
            return resp.json()

        # === Step 3: Fetch first page & detect total pages ===
        first_page = fetch_page(1)

        # IMPORTANT: If your API uses a different field name, change this.
        total_pages = (
            first_page.get("TotalPages")
            or first_page.get("totalPages")
            or first_page.get("Meta", {}).get("TotalPages")
            or 1
        )

        logging.info(f"Total pages detected: {total_pages}")

        # === Step 4: Loop through ALL pages ===
        all_data = []
        all_data.append(first_page)

        for p in range(2, total_pages + 1):
            logging.info(f"Fetching page {p}/{total_pages}...")
            page_data = fetch_page(p)
            all_data.append(page_data)

        # === Step 5: Create Word document ===
        doc = Document()
        doc.add_heading("API Data Export", level=1)
        doc.add_paragraph(f"Fetched from: {data_url}")
        doc.add_paragraph(f"From: {from_param}")
        doc.add_paragraph(f"To: {to_param}")
        doc.add_paragraph(f"Total Pages: {total_pages}")
        doc.add_paragraph("")

        # Recursive writer
        def write_json_to_doc(d, indent=0):
            if isinstance(d, dict):
                for key, value in d.items():
                    if isinstance(value, (dict, list)):
                        doc.add_paragraph(" " * indent + f"{key}:", style="List Bullet")
                        write_json_to_doc(value, indent + 2)
                    else:
                        doc.add_paragraph(" " * indent + f"{key}: {value}")
            elif isinstance(d, list):
                for item in d:
                    write_json_to_doc(item, indent + 2)
            else:
                doc.add_paragraph(str(d))

        # Write combined data
        doc.add_heading("Combined API Results", level=2)
        write_json_to_doc(all_data)

        # === Step 6: Return file ===
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        filename = f"api_data_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"

        return func.HttpResponse(
            body=stream.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
            status_code=200,
        )

    except Exception as e:
        logging.error(f"Error: {e}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            status_code=500,
        )
